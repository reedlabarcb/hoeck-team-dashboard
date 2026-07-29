#!/usr/bin/env python3
"""
Word (.docx) + Excel (.xlsx) text extractor for the Hoeck Team Dashboard.

Invoked as a subprocess by lib/external/box/text-extractor.ts, exactly like
scripts/python/pdf_extract_text.py. Reads a single office file (downloaded locally
by the TS worker from Box) and emits JSON to stdout.

CONTRACT: identical to pdf_extract_text.py — same keys, same status enum, same exit
codes — so the TS worker's status->extraction_status mapping and DB writes are
unchanged. Do not add or rename statuses here without changing that mapping AND the
extraction_status Postgres enum.

Usage:
    python office_extract_text.py --file-path /tmp/in.docx
    python office_extract_text.py --file-path /tmp/in.xlsx

Returns (stdout, always JSON):
    {
      "status":           "ok" | "scanned" | "too_large" | "error",
      "text":             str | null,
      "page_count":       int,
      "character_count":  int,
      "extraction_method": "python-docx" | "openpyxl" | "office",
      "warnings":         [str, ...],
      "error":            str       # status="error" only
    }

Status meanings (same DB mapping as the PDF path):
    "ok"         — text extracted (caller writes extracted_text + 'extracted')
    "scanned"    — opened fine but yielded NO text at all: an empty document, or a
                   sheet holding only formatting/images. Reuses the PDF path's
                   "opened but no text" status so the caller writes
                   'skipped_scanned' — no new enum value, no schema change.
                   (The label is PDF-flavored; the meaning here is "no extractable
                   text", which is the same thing the caller needs to know.)
    "too_large"  — file size > MAX_FILE_BYTES, or a workbook past MAX_CELLS.
                   Caller writes 'skipped_too_large'.
    "error"      — extraction threw. Caller writes 'failed' + stores `error`.

page_count semantics per type (the caller stores it in box_folder_index.page_count,
and maps 0 -> NULL):
    .xlsx -> number of worksheets
    .docx -> 0 (a .docx has no page count without rendering it; reporting a fake
             number would be worse than NULL)

Exit codes:
    0 = success (status in {"ok", "scanned", "too_large"})
    1 = extraction error (status == "error")
    2 = argument error
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from typing import Any

# ----- Tunables (mirror pdf_extract_text.py where they overlap) -----

# Hard ceiling on file size we'll attempt to read. Same 50 MB as the PDF path,
# for the same reason: avoid OOM on Railway's small worker dynos.
MAX_FILE_BYTES: int = 50 * 1024 * 1024  # 50 MB

# Workbook cell ceiling. A 50 MB xlsx can still expand to tens of millions of cells;
# stringifying them all would blow memory even though the FILE passed the size guard.
# Past this we bail with "too_large" rather than risk the worker.
MAX_CELLS: int = 500_000

# Below this many characters the document has effectively no text -> "scanned".
# 1 (not the PDF's 100): a one-cell spreadsheet or a one-line memo is legitimately
# tiny and IS searchable content, so only a genuinely empty result is skipped.
MIN_CHARS_FOR_CONTENT: int = 1

SUPPORTED_EXTENSIONS = ('.docx', '.xlsx')


def _payload(
    status: str,
    *,
    text: str | None = None,
    page_count: int = 0,
    character_count: int = 0,
    method: str = 'office',
    warnings: list[str] | None = None,
    error: str | None = None,
) -> dict[str, Any]:
    """Build the response dict. Key set is fixed by the shared contract."""
    out: dict[str, Any] = {
        'status': status,
        'text': text,
        'page_count': page_count,
        'character_count': character_count,
        'extraction_method': method,
        'warnings': warnings or [],
    }
    if error is not None:
        out['error'] = error
    return out


def _emit(payload: dict[str, Any], exit_code: int) -> None:
    """Write JSON to stdout and exit. Same shape as pdf_extract_text.py."""
    json.dump(payload, sys.stdout, default=str, indent=None)
    sys.stdout.write('\n')
    sys.exit(exit_code)


def _clean(text: str) -> str:
    """Strip NUL (0x00) — Postgres `text` columns reject it outright. Same strip the
    PDF extractor does at the source; text-extractor.ts scrubs again before the write."""
    return text.replace('\x00', '')


# ----------------- .docx -----------------

def _extract_docx(file_path: str) -> tuple[dict[str, Any], int]:
    """Paragraphs AND table cells. CRE documents carry real content in tables
    (rent schedules, option dates, premises), and python-docx does NOT include
    table text in document.paragraphs — skipping tables would silently drop it."""
    try:
        import docx  # noqa: PLC0415  (python-docx)
    except ImportError as e:
        return (
            _payload(
                'error',
                method='python-docx',
                error=(
                    f'python-docx not installed ({e}). It ships via the pip line in '
                    'nixpacks.toml on Railway; locally `pip install python-docx`.'
                ),
            ),
            1,
        )

    try:
        document = docx.Document(file_path)
        parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        # Table cell text, row by row — tab-joined per row so columns stay distinguishable.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append('\t'.join(cells))
    except Exception as e:  # noqa: BLE001 — corrupt docx must surface as JSON, not a crash
        return (
            _payload('error', method='python-docx', error=f'python-docx raised {type(e).__name__}: {e}'),
            1,
        )

    full_text = _clean('\n'.join(parts))
    count = len(full_text)

    if count < MIN_CHARS_FOR_CONTENT:
        return (
            _payload(
                'scanned',
                method='python-docx',
                character_count=count,
                warnings=[
                    'Document opened but yielded no text (no paragraphs or table cells with '
                    'content). Likely an image-only or empty .docx.'
                ],
            ),
            0,
        )

    # page_count stays 0 -> the caller stores NULL (a .docx has no page count unrendered).
    return (_payload('ok', text=full_text, character_count=count, method='python-docx'), 0)


# ----------------- .xlsx -----------------

def _extract_xlsx(file_path: str) -> tuple[dict[str, Any], int]:
    """ALL sheets, read_only + data_only.

    data_only=True is essential: without it openpyxl hands back FORMULA STRINGS
    ("=SUM(B2:B9)") instead of the cached computed values, which would index
    gibberish. read_only=True keeps memory flat on large workbooks.
    """
    try:
        from openpyxl import load_workbook  # noqa: PLC0415
    except ImportError as e:
        return (
            _payload(
                'error',
                method='openpyxl',
                error=(
                    f'openpyxl not installed ({e}). It ships via the pip line in '
                    'nixpacks.toml on Railway; locally `pip install openpyxl`.'
                ),
            ),
            1,
        )

    workbook = None
    try:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
        sheet_count = len(workbook.sheetnames)
        parts: list[str] = []
        cells_seen = 0
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            # Sheet NAME is content too ("2027 Renewals", "Suite 400") — index it.
            parts.append(str(sheet_name))
            for row in sheet.iter_rows(values_only=True):
                values: list[str] = []
                for value in row:
                    cells_seen += 1
                    if value is None or value == '':
                        continue
                    values.append(str(value))
                if cells_seen > MAX_CELLS:
                    return (
                        _payload(
                            'too_large',
                            page_count=sheet_count,
                            method='openpyxl',
                            warnings=[
                                f'Workbook exceeds MAX_CELLS ({MAX_CELLS:,}) — stopped after '
                                f'{cells_seen:,} cells. Skipped without indexing.'
                            ],
                        ),
                        0,
                    )
                if values:
                    parts.append('\t'.join(values))
    except Exception as e:  # noqa: BLE001 — corrupt xlsx must surface as JSON, not a crash
        return (
            _payload('error', method='openpyxl', error=f'openpyxl raised {type(e).__name__}: {e}'),
            1,
        )
    finally:
        # read_only workbooks hold an open zip handle; close it or /tmp cleanup can fail on Windows.
        if workbook is not None:
            try:
                workbook.close()
            except Exception:  # noqa: BLE001
                pass

    full_text = _clean('\n'.join(parts))
    count = len(full_text)

    if count < MIN_CHARS_FOR_CONTENT:
        return (
            _payload(
                'scanned',
                page_count=sheet_count,
                character_count=count,
                method='openpyxl',
                warnings=['Workbook opened but every cell was empty.'],
            ),
            0,
        )

    return (
        _payload('ok', text=full_text, page_count=sheet_count, character_count=count, method='openpyxl'),
        0,
    )


# ----------------- shared entry -----------------

def extract(file_path: str) -> tuple[dict[str, Any], int]:
    """Inspect + extract, dispatching on extension. Returns (payload, exit_code).

    Pure function so pytest can call it directly without argparse + sys.exit —
    same structure as pdf_extract_text.extract().
    """
    if not os.path.isfile(file_path):
        return (_payload('error', error=f'File not found: {file_path}'), 1)

    lower = file_path.lower()
    if not lower.endswith(SUPPORTED_EXTENSIONS):
        return (
            _payload(
                'error',
                error=(
                    f'Unsupported extension for {file_path}. This extractor handles '
                    f'{", ".join(SUPPORTED_EXTENSIONS)}; PDFs go to pdf_extract_text.py.'
                ),
            ),
            1,
        )

    # Size guard FIRST — before opening — same order as the PDF path.
    try:
        size = os.path.getsize(file_path)
    except OSError as e:
        return (_payload('error', error=f'Could not stat file: {type(e).__name__}: {e}'), 1)

    if size > MAX_FILE_BYTES:
        return (
            _payload(
                'too_large',
                warnings=[
                    f'File size {size:,} bytes exceeds MAX_FILE_BYTES '
                    f'({MAX_FILE_BYTES:,} bytes). Skipped without attempting extraction.'
                ],
            ),
            0,
        )

    return _extract_docx(file_path) if lower.endswith('.docx') else _extract_xlsx(file_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description='Extract text from a single .docx or .xlsx file.')
    parser.add_argument('--file-path', required=True, help='Local path to the .docx / .xlsx file.')
    args = parser.parse_args(argv)

    payload, exit_code = extract(args.file_path)
    _emit(payload, exit_code)
    return exit_code  # unreachable — _emit calls sys.exit, but satisfies type checker


if __name__ == '__main__':
    sys.exit(main())
